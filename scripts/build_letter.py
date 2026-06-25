"""
Build Globeleq_Application_Letter.docx
Hiring-manager cover letter for Associate BI Specialist role.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = r"C:\Users\Anthony.DESKTOP-ES5HL78\Documents\Globeleq_Energy_Intelligence_Platform\Globeleq_Application_Letter.docx"

GREEN  = RGBColor(0x14, 0x44, 0x3B)   # #14443B  forest green
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LTGRN  = RGBColor(0xEA, 0xF6, 0xF3)   # #EAF6F3  light green tint
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
AMBER  = RGBColor(0xF7, 0x94, 0x1D)

# ── helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def set_cell_borders(cell, color="C8E8E1"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def cell_para(cell, text, bold=False, italic=False,
              color=BLACK, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name  = "Calibri"
    return p

def add_heading(doc, text, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold  = True
    run.font.size  = Pt(size)
    run.font.color.rgb = GREEN
    run.font.name  = "Calibri"
    return p

def add_body(doc, text, italic=False, color=BLACK, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.italic = italic
    run.font.size  = Pt(11)
    run.font.color.rgb = color
    run.font.name  = "Calibri"
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.color.rgb = BLACK
    run.font.name  = "Calibri"
    return p

def add_divider(doc, color="14443B"):
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.name = "Calibri"
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

# ── document setup ────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.5)
add_page_number(section)

# default paragraph style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── LETTERHEAD ────────────────────────────────────────────────────────────────
# Green top bar (paragraph with bottom border + green shading trick via table-free approach)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
pPr = p._p.get_or_add_pPr()
# green paragraph shading
shd = OxmlElement("w:shd")
shd.set(qn("w:val"),   "clear")
shd.set(qn("w:color"), "auto")
shd.set(qn("w:fill"),  "14443B")
pPr.append(shd)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
r1 = p.add_run("  ANTHONY APOLLIS")
r1.font.name  = "Calibri"
r1.font.size  = Pt(16)
r1.font.bold  = True
r1.font.color.rgb = WHITE
r2 = p.add_run("   |   Data Engineer & ML Practitioner")
r2.font.name  = "Calibri"
r2.font.size  = Pt(11)
r2.font.color.rgb = LTGRN

# contact line
pc = doc.add_paragraph()
pc.paragraph_format.space_before = Pt(2)
pc.paragraph_format.space_after  = Pt(8)
pPrc = pc._p.get_or_add_pPr()
shd2 = OxmlElement("w:shd")
shd2.set(qn("w:val"),   "clear")
shd2.set(qn("w:color"), "auto")
shd2.set(qn("w:fill"),  "14443B")
pPrc.append(shd2)
rc = pc.add_run("  anthony.apollis@gmail.com   |   github.com/anthonyapollis   |   Cape Town, South Africa")
rc.font.name  = "Calibri"
rc.font.size  = Pt(9.5)
rc.font.color.rgb = LTGRN

add_divider(doc, "F7941D")   # amber rule under letterhead

# ── DATE + RECIPIENT ─────────────────────────────────────────────────────────
p_date = doc.add_paragraph()
p_date.paragraph_format.space_before = Pt(10)
p_date.paragraph_format.space_after  = Pt(0)
rd = p_date.add_run("25 June 2026")
rd.font.size = Pt(10); rd.font.name = "Calibri"; rd.font.color.rgb = RGBColor(0x60,0x60,0x60)

for line in ["The Hiring Manager", "Globeleq Management Services",
             "Associate BI Specialist — South Africa"]:
    pr = doc.add_paragraph()
    pr.paragraph_format.space_before = Pt(0)
    pr.paragraph_format.space_after  = Pt(0)
    rr = pr.add_run(line)
    rr.font.size = Pt(11)
    rr.font.name = "Calibri"
    rr.font.color.rgb = BLACK
    if line == "Associate BI Specialist — South Africa":
        rr.italic = True
        rr.font.color.rgb = RGBColor(0x40,0x40,0x40)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── SUBJECT ───────────────────────────────────────────────────────────────────
ps = doc.add_paragraph()
ps.paragraph_format.space_before = Pt(2)
ps.paragraph_format.space_after  = Pt(10)
rs1 = ps.add_run("Subject: ")
rs1.bold = True; rs1.font.size = Pt(11); rs1.font.name = "Calibri"; rs1.font.color.rgb = GREEN
rs2 = ps.add_run("Application — Associate BI Specialist | Globeleq Energy Intelligence Platform Portfolio")
rs2.bold = True; rs2.font.size = Pt(11); rs2.font.name = "Calibri"; rs2.font.color.rgb = BLACK

# ── OPENING ───────────────────────────────────────────────────────────────────
add_body(doc,
    "Dear Hiring Manager,",
    space_after=8)
add_body(doc,
    "Please find attached my application for the Associate BI Specialist position at Globeleq. "
    "To demonstrate applied competency across the Globeleq operating environment — African IPP power plants, "
    "SCADA telemetry, PPA commercial structures, maintenance reliability, and ESG reporting — I built a "
    "full end-to-end data engineering and ML intelligence platform using Globeleq's own portfolio as the domain. "
    "The full source code is available for review at: "
    "https://github.com/anthonyapollis/globeleq-energy-intelligence-platform",
    space_after=10)

# ── SECTION: WHAT I BUILT ─────────────────────────────────────────────────────
add_heading(doc, "What I Built")
add_body(doc,
    "A production-grade Azure Databricks data platform ingesting synthetic SCADA telemetry from "
    "17 operating plants (1,794 MW) and 2 under-construction plants (485 MW) across 7 African countries, "
    "covering the period 2020 to 2024.",
    space_after=8)

deliverables = [
    ("Azure Databricks Medallion Pipeline",
     "Bronze ingest → Silver cleanse (imputation, outlier detection) → Gold KPIs across 4 PySpark notebooks"),
    ("Azure Data Factory Pipeline",
     "ForEach activity (batchCount=4, 8 fact tables), IfCondition ML gate, WebActivity for Teams alert + "
     "Power BI dataset refresh, daily 02:00 SAST trigger"),
    ("SQL Server Data Warehouse",
     "22 tables across dw + bi schemas; 12 core + 10 V2 tables including FactMeterReadingHourly "
     "(raw/clean/quality split) and FactDataQualityEvent (immutable audit log)"),
    ("SCADA Dataset",
     "3,024,807 rows total; 2,981,664 rows 15-min telemetry (17 plants × 5 years × 96 intervals/day)"),
    ("Dirty Data Layer",
     "35 realistic data quality patterns injected: 57,176 null sensor dropouts, 4,472 fault-code 999.9 readings, "
     "5,366 out-of-range power values, 8,944 duplicate rows — Silver notebook cleans all with full audit trail"),
    ("Power BI",
     "111 DAX measures across 9 report pages: Executive Overview, Control Room, Forecast Operations, "
     "Model Scorecard, Meter Data Quality, Weather-Normalised Performance, Maintenance & Reliability, "
     "Construction Portfolio, ESG & HSE"),
    ("Excel Report",
     "8-sheet workbook with BarChart, LineChart, PieChart, conditional formatting — attached"),
    ("HTML Ebook",
     "8-chapter technical walkthrough with Chart.js visualisations — attached (open in any browser)"),
]

col_w = [Cm(5.5), Cm(12.1)]
tbl1 = doc.add_table(rows=1 + len(deliverables), cols=2)
tbl1.style = "Table Grid"
tbl1.autofit = False
tbl1.columns[0].width = col_w[0]
tbl1.columns[1].width = col_w[1]

# header
hdr = tbl1.rows[0]
for i, txt in enumerate(("Deliverable", "Detail")):
    c = hdr.cells[i]
    c.width = col_w[i]
    set_cell_bg(c, GREEN)
    set_cell_borders(c, "14443B")
    cell_para(c, txt, bold=True, color=WHITE, size=10)

for ri, (deliv, detail) in enumerate(deliverables):
    row = tbl1.rows[ri + 1]
    bg  = LTGRN if ri % 2 == 0 else WHITE
    for ci, txt in enumerate([deliv, detail]):
        c = row.cells[ci]
        c.width = col_w[ci]
        set_cell_bg(c, bg)
        set_cell_borders(c, "C8E8E1")
        cell_para(c, txt, bold=(ci == 0), size=10)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── SECTION: ML MODELS ────────────────────────────────────────────────────────
add_heading(doc, "Machine Learning Results")
add_body(doc, "8 models registered in MLflow across two tiers:", space_after=6)

ml_rows = [
    ("Energy Yield Forecaster",           "XGBoost Regressor",            "R²=0.94, MAE=18 MWh"),
    ("Forced Outage Predictor",           "LightGBM Classifier",          "AUC=0.82, AP=0.61 (scale_pos_weight handles 11× class imbalance; threshold at 0.35 for recall)"),
    ("Maintenance Cost Estimator",        "Random Forest",                 "R²=0.89, OOB score=0.87"),
    ("Curtailment Anomaly Detector",      "Isolation Forest",              "5% anomaly rate, contamination=0.05"),
    ("Portfolio Revenue Forecaster",      "LightGBM Regressor",           "R²=0.93, MAPE=3.1%"),
    ("Weekly Seasonal Profile",           "Seasonal Decomposition",        "Baseline forecasting model"),
    ("Robust Dynamic Regression",         "MM-Estimator Regression",       "Challenger model (outlier-robust training)"),
    ("Technology-Aware Weather GBM ★", "Gradient Boosting",           "SELECTED — lowest RMSE; tech-specific features: GHI (solar), wind speed + air density (wind), dispatch + heat rate (gas)"),
]

ml_col_w = [Cm(5.0), Cm(4.5), Cm(8.1)]
tbl2 = doc.add_table(rows=1 + len(ml_rows), cols=3)
tbl2.style = "Table Grid"
tbl2.autofit = False
tbl2.columns[0].width = ml_col_w[0]
tbl2.columns[1].width = ml_col_w[1]
tbl2.columns[2].width = ml_col_w[2]

for i, txt in enumerate(("Model", "Algorithm", "Result")):
    c = tbl2.rows[0].cells[i]
    c.width = ml_col_w[i]
    set_cell_bg(c, GREEN)
    set_cell_borders(c, "14443B")
    cell_para(c, txt, bold=True, color=WHITE, size=10)

for ri, (model, algo, result) in enumerate(ml_rows):
    row = tbl2.rows[ri + 1]
    bg  = LTGRN if ri % 2 == 0 else WHITE
    is_selected = "★" in model
    for ci, txt in enumerate([model, algo, result]):
        c = row.cells[ci]
        c.width = ml_col_w[ci]
        set_cell_bg(c, bg)
        set_cell_borders(c, "C8E8E1")
        cell_para(c, txt, bold=is_selected, size=10)

doc.add_paragraph().paragraph_format.space_after = Pt(2)
add_body(doc,
    "Model selection uses lowest validation RMSE with absolute total forecast bias as tie-breaker. "
    "MAPE is stored and reported but is not used as the primary selection metric — it is unreliable "
    "for solar night-time and other near-zero readings.",
    italic=True, color=RGBColor(0x50,0x50,0x50), space_after=10)

# ── SECTION: BUSINESS INSIGHTS ────────────────────────────────────────────────
add_heading(doc, "Key Business Insights from the Data")

insights = [
    ("Natural Gas concentration risk",
     "Azito Power (713 MW, Côte d’Ivoire) generates 24.4 TWh — 52% of the entire portfolio. "
     "Gas plants run at 69.3% availability vs 92%+ for solar and wind. A 1% availability improvement "
     "on Azito alone is worth an estimated R4.3M in additional annual revenue at current PPA tariff rates."),
    ("Forced outage trend",
     "Outages declined 2020→2023 (129 to 96) then reversed to 128 in 2024 — a classic maintenance "
     "backlog signal. The Forced Outage Predictor identifies the 7-day precursor window, enabling "
     "pre-emptive intervention before generation loss occurs."),
    ("ESG positioning",
     "5.02 Mt CO₂e avoided over 5 years. Portfolio Scope 1 intensity is 406 tCO₂e/GWh from gas. "
     "The Menengai Geothermal plant (35 MW, Kenya) brings this down meaningfully — geothermal runs "
     "at ~50 tCO₂e/GWh — directly relevant to BII and Norfund sustainability KPIs."),
    ("Settlement collection quality",
     "Revenue averaging R139M/year is flat (PPA structure), but collection percentage anomalies in "
     "the energy sales data flag a real cash-flow risk that the Meter Data Quality dashboard surfaces "
     "automatically, with every event logged in FactDataQualityEvent for audit."),
]

for label, body in insights:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rl = p.add_run(label + ": ")
    rl.bold = True
    rl.font.size = Pt(11); rl.font.name = "Calibri"; rl.font.color.rgb = GREEN
    rb = p.add_run(body)
    rb.font.size = Pt(11); rb.font.name = "Calibri"; rb.font.color.rgb = BLACK

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── SECTION: ATTACHMENTS ──────────────────────────────────────────────────────
add_heading(doc, "Attachments")

attachments = [
    "Globeleq_Energy_Intelligence_Report.xlsx — 8-sheet Excel workbook: Portfolio KPIs, Plant Performance, Technology Mix, Revenue Analysis, Reliability, HSE Dashboard, ML Results",
    "globeleq_energy_intelligence_ebook.html — 8-chapter technical ebook with Chart.js visualisations (open in any browser)",
    "Azure Data Factory pipeline screenshot — pl_globeleq_daily_energy_intelligence (ForEach, IfCondition, WebActivity)",
    "Azure Databricks architecture diagram — Medallion Bronze → Silver → Gold with Delta Lake",
    "Power BI report page mockups — 9 pages, 111 measures",
]
for att in attachments:
    add_bullet(doc, att)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

p_git = doc.add_paragraph()
p_git.paragraph_format.space_before = Pt(2)
p_git.paragraph_format.space_after  = Pt(10)
p_git.alignment = WD_ALIGN_PARAGRAPH.LEFT
rg1 = p_git.add_run("Full source code, notebooks, SQL DDL, DAX measures, ADF pipeline JSON: ")
rg1.font.size = Pt(11); rg1.font.name = "Calibri"; rg1.font.color.rgb = BLACK
rg2 = p_git.add_run("https://github.com/anthonyapollis/globeleq-energy-intelligence-platform")
rg2.font.size = Pt(11); rg2.font.name = "Calibri"; rg2.font.color.rgb = GREEN
rg2.bold = True

add_divider(doc, "14443B")

# ── CLOSING ────────────────────────────────────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_body(doc,
    "I am available for a technical interview at short notice ahead of the 7 July 2026 closing date. "
    "I am based in Cape Town and comfortable with the hybrid working arrangement described. "
    "I look forward to discussing how this work demonstrates readiness for the BI Specialist role.",
    space_after=14)

for line in ["Yours sincerely,", "", "Anthony Apollis",
             "anthony.apollis@gmail.com", "github.com/anthonyapollis"]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1 if line else 6)
    r = p.add_run(line)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    if line in ("Anthony Apollis",):
        r.bold = True
        r.font.color.rgb = GREEN
    elif line in ("anthony.apollis@gmail.com", "github.com/anthonyapollis"):
        r.font.color.rgb = GREEN
    else:
        r.font.color.rgb = BLACK

# ── SAVE ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
